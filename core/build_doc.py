import os
from docx import Document
from docx.shared import RGBColor, Pt


def build_doc(songs, filename, template=None):
    filename = os.path.abspath(filename)
    folder = os.path.dirname(filename)
    print(folder)
    
    doc = Document()
    
    doc.styles['Normal'].font.name = 'Raleway'
    doc.styles['Normal'].font.size = Pt(10)
    doc.styles['Heading 1'].font.name = 'Raleway'
    doc.styles['Heading 1'].font.size = Pt(14)
    doc.styles['Heading 1'].font.color.rgb = RGBColor(0, 0, 0)
    doc.styles['Heading 2'].font.name = 'Raleway'
    doc.styles['Heading 2'].font.size = Pt(12)
    doc.styles['Heading 2'].font.color.rgb = RGBColor(0, 0, 0)

    for song in songs:
        doc.add_heading(song['title'], 1)
        doc.add_heading(f"{song['name']} - {song['artist']}", 2)
        doc.add_paragraph(song['lyrics'])

    if not '.docx' in filename:
        filename += '.docx'

    doc.save(filename)
    
    
if __name__ == '__main__':
    build_doc(..., '')